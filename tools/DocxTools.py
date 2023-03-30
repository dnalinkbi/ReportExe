import docx
import tkinter as tk
import tkinter.ttk as ttk

def Make_UI_Docx():
    # Create the window
    root = tk.Tk()
    root.title("Illumina Docx Report Maker")
    root.geometry("600x600+600+200")

    # Create the labels and text entry boxes
    label_1 = tk.Label(root, text="Enter the title of your document:")
    entry_1 = tk.Entry(root)
    label_2 = tk.Label(root, text="Enter the body of your document:")
    entry_2 = tk.Text(root, height=10, width=50)

    # Place the widgets in the window
    label_1.grid(row=0, column=0, sticky=tk.E)
    entry_1.grid(row=0, column=1)
    label_2.grid(row=1, column=0, sticky=tk.E)
    entry_2.grid(row=1, column=1)

    # Define the list of variables Libaries 
    Var_Platform = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']

    # Library of ComboBox widgets
    Combobox_label = tk.Label(root, text="Select a library")
    Combobox_label.grid(row=2, column=0, sticky=tk.E)
    Combobox_Platform = ttk.Combobox(root, height=5, values=Var_Platform)
    Combobox_Platform.set("Library")
    Combobox_Platform.grid(row=2, column=1)

    # Create the function to write the document
    def create_document():
        document = docx.Document()
        document.add_heading(entry_1.get(), 0)
        document.add_paragraph(entry_2.get("1.0", tk.END))
        document.save('my_document.docx')

    # Create the button
    button = tk.Button(root, text="Create Document", command=create_document)
    button.grid(row=3, columnspan=2)

    # Run the window
    root.mainloop()