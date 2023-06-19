
from docx import Document
from docx.shared import Inches
from docx.shared import Mm
from docx.shared import Cm
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.text.font import Font
from docx.text.parfmt import ParagraphFormat
from docx.shared import Length
from docx.oxml import parse_xml
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
from pandas import DataFrame
import subprocess
from pandas import read_csv
import sys
import numpy as np
# from HL.Add_HL import add_hyperlink as HL
#from matplotlib import pyplot as plt

ServiceID=sys.argv[1]


##find institution, client Korrea name##
ID = ServiceID.split('_')
institution = ID[1]
client = ID[2]
input =  pd.read_csv('/ess/dlstibm/Workspace/workspace.ckw/DB/salesforce_client_Eng.txt', delimiter='\t', encoding = "euc-kr")
institutionkor = input[input[u'단축코드'] == "%s"%institution].iloc[0,1]
clientkor = input[input[u'단축코드'] == "%s"%client].iloc[0,1]

#################################################################
#Title#---------------------------------------------------------#
#################################################################


document = Document()
section = document.sections[0]
section.left_margin = Mm(0.0)
section.right_margin = Mm(0.0)
section.top_margin = Mm(0.0)
section.bottom_margin = Mm(0.0)
date = subprocess.check_output('date "+%Y.%m.%d"', shell=True)
date = date.decode()
#print "%s"%date

headtailpng='/ess/dlstibm/Workspace/workspace.ckw/tool/docx/dnalink_logo.png'

section = document.sections[0]


headers = section.header
head = headers.add_paragraph()
head_format = head.paragraph_format
head2 = head.add_run()
head2.add_picture(headtailpng, width=Inches(2))
head_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

#document.add_picture(headtailpng, width=Inches(9.200))
#document.add_heading('Document Title', 0)
#document.add_heading('./title.jpeg')
#document.add_picture('./headtail.png', width=Inches(9.200))

A = document.add_paragraph()
H = A.add_run('\t\tDNA Link')
H.font.name = 'Century'
H.font.size = Pt(40)
H.font.bold = True
E = A.add_run('\n\t\tGenomic Service')
E.font.name = 'Century'
E.font.size = Pt(40)
E.font.bold = True
D = A.add_run('\n\t\tSequencing Report')
D.font.name = 'Century'
D.font.size = Pt(40)
D.font.bold = True


document.add_paragraph('\t\t\t\t')
NGS = document.add_paragraph()
Next = NGS.add_run('\t\t\t\tNext Generation Seqencing')
Next.font.name = 'Microsoft Sans Serif'
Next.font.size = Pt(20)
Next.font.color.rgb = RGBColor(183, 183, 183)
Name = document.add_paragraph()
Name_format = Name.paragraph_format
#Name_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
send = Name.add_run(u'\t\t\t\t\t\t\t\t\t\t\t%s %s 선생님 귀하'%(institutionkor,clientkor))
send.font.name = 'Myriad Pro'
send.font.size = Pt(10)
send.font.bold = True

title4 = document.add_paragraph()
title4_format = title4.paragraph_format
center = title4.add_run()
center.add_picture('/ess/dlstibm/Workspace/workspace.ckw/tool/docx/title4.png', width=Cm(22.14), height=Cm(11.85))


section = document.sections[0]
footers = section.footer
foot = footers.add_paragraph()
foot2 = foot.add_run('\tDNA Link Genomic Service Sequencing Report\t\t\t\t\t\t\t\t %s'%date)
#foot2.add_picture(headtailpng, width=Inches(10.100))

#################################################################
#Work flow------------------------------------------------------#
#################################################################

new_section = document.add_section(1)

workflow = document.add_paragraph()
wf = workflow.add_run('\tNovaSeq 6000 Workflow Overview')
wf.font.name = 'Calibri'
wf.font.size = Pt(18)
wf.font.bold = True

method = document.add_paragraph()
method_format = method.paragraph_format
center = method.add_run()
center.add_picture('/ess/dlstibm/Workspace/workspace.ckw/tool/docx/method.png', width=Inches(7.100))
method_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

methodmanual = document.add_paragraph()
mt = methodmanual.add_run('\n\tNovaSeq6000Sequencing')
mt.font.name = 'Calibri'
mt.font.size = Pt(11)
mt.font.bold = True
methodmanual2 = document.add_paragraph()
mt2 = methodmanual2.add_run(u'\tQC를 통과한 DNA는 illumina의 DNA sample prep kit를 이용하여 library로 제작됩니다. 완성된 DNA library는 Illumina NovaSeq 6000을 \n\t이용하여 paired end read 101bp로 sequencing 한다. 준비된 library를 Cluster Generation 장치인 cBot에 장착된 flow cell에 loading하고 \n\tbridge amplification 방법을 이용하여 template를 증폭한 후, flow cell을 옮겨, sequencing 반응을 시작합니다. Sequencing 반응은 한 cycle에 \n\t서로 다른 형광 dye를 가지는 dNTP를 한 base씩 합성해 나가면서 이용된 nucleotide의 형광을 측정하여 sequence를 분석하게 됩니다.')
mt2.font.name = 'Myriad Pro'
mt2.font.size = Pt(9)

#################################################################
#Experiment process---------------------------------------------#
#################################################################

new_section = document.add_section(1)
process = document.add_paragraph()
process2 = process.add_run('\t\tExperiment process & the time required')
process2.font.name = 'Calibri'
process2.font.size = Pt(18)
process2.font.bold = True
process.add_run('\n')
process.add_run('\n')
process3 = process.add_run('\t\t<Workflow>')
process3.font.name = 'Calibri'
process3.font.size = Pt(16)
process3.font.bold = True
method2 = document.add_paragraph()
method2_format = method2.paragraph_format
center = method2.add_run()
center.add_picture('/ess/dlstibm/Workspace/workspace.ckw/tool/docx/method2.jpeg', width=Inches(7.100))
method2_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

analysis = document.add_paragraph()
analysis2 = analysis.add_run(u'\n\n\t\t<분석 세부 내용>')
analysis2.font.name = 'Myriad Pro'
analysis2.font.size = Pt(16)
analysis2.font.bold = True
analysis3 = analysis.add_run(u'\n\n\t\t1) DNA Extraction\n\n\t\t2) Sample Quality Control\n\n\t\t3) Library 제작\n\n\t\t4) Sequencing: Illumina Novaseq 6000이용\n\n\t\t5) Bioinformatics Analysis - Raw data filtering\n\n\t\t6) 최종 결과 보고서')
analysis3.font.name = 'Myriad Pro'
analysis3.font.size = Pt(14)

#################################################################
#Sequencing & Sample Info---------------------------------------#
#################################################################

new_section = document.add_section(1)
SequenceReport = document.add_paragraph()
SequenceReport2 = SequenceReport.add_run('\t\tSequencing Report')
SequenceReport.add_run('\n')
SequenceReport2.font.name = 'Calibri'
SequenceReport2.font.size = Pt(18)
SequenceReport2.font.bold = True
infomation = document.add_paragraph()
infomation2 = infomation.add_run('\t\t Service Information')
infomation2.font.name = 'Calibri'
infomation2.font.size = Pt(14)
infomation2.font.bold = True
infomation3 = infomation.add_run('\n\t\t The table below contains the service information related to the samples you have submitted.')
infomation3.font.name = 'Calibri'
infomation3.font.size = Pt(12)
infomation.add_run('\n')



## pandas test
try:
	infofile = pd.read_csv("./%s_docx" %ServiceID, sep="\t", header=None, encoding='utf-8')
except:
	infofile = pd.read_csv("./%s_docx" %ServiceID, sep="\t", header=None, encoding='euc-kr')
#print infofile


rows = len(infofile.iloc[:,0])
cols = len(infofile.iloc[0,:])
#print rows, cols

table = document.add_table(rows = rows, cols = cols)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = document.styles['Medium Grid 2 Accent 3']

for i,j in enumerate(infofile.iloc[:,0]):
	#print i,j
	for x,y in enumerate(infofile.iloc[i,:]):
		#print x, y
		table.rows[i].cells[x].paragraphs[0].add_run('%s'%y).font.size = Pt(12)
		#table.rows[i].cells[1].paragraphs[0].add_run('%s'%y)
		table.rows[i].cells[x].width = Cm(8.16)
		table.rows[i].cells[x].height = Cm(5)
		#shading_elm = parse_xml(r'<w:shd {} w:fill="93DAFF"/>'.format(nsdecls('w')))
		#table.rows[i].cells[0]._tc.get_or_add_tcPr().append(shading_elm)
		#shading_elm2 = parse_xml(r'<w:shd {} w:fill="C0FFFF"/>'.format(nsdecls('w')))
		#able.rows[i].cells[1]._tc.get_or_add_tcPr().append(shading_elm2)
		#table.rows[i].cells[0].text_frame.paragraphs[0].run.font.size = Pt(12)
		#table.rows[i].cells[0].text_frame.paragraphs[0].run.font.color.rgb = RGBColor(255, 255, 255)

blank = document.add_paragraph()		
blank.add_run('\n')

#Sample Info ##############

seqinof = document.add_paragraph()
seqinof2 = seqinof.add_run('\t\t Sequencing Information')
seqinof2.font.name = 'Calibri'
seqinof2.font.size = Pt(14)
seqinof2.font.bold = True
seqinof.add_run('\n')
seqinof3 = seqinof.add_run('\n\t\t The table below lists all the samples processed in this project and their specifications according\n\t\t  to the sample submission form.')
seqinof3.font.name = 'Calibri'
seqinof3.font.size = Pt(12)
seqinof.add_run('\n')


infofile2 = pd.read_csv("./%s_sample" %ServiceID, sep="\t", header=None, engine="python")
rows2 = len(infofile2.iloc[:,0])
cols2 = len(infofile2.iloc[0,:])
table2 = document.add_table(rows = rows2, cols = cols2)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
#table2.style = document.styles['Light List Accent 5']
table2.style = document.styles['Medium Grid 2 Accent 3']

for i,j in enumerate(infofile2.iloc[:,0]):
	for x,y in enumerate(infofile2.iloc[i,:]):
		table2.rows[i].cells[x].height = Cm(5)
		

for i, j in enumerate(infofile2.iloc[:,0]):
		table2.rows[i].cells[0].paragraphs[0].add_run('%s'%j).font.size = Pt(12)
		table2.rows[i].cells[0].width = Cm(3.0)
		
for a, b in enumerate(infofile2.iloc[:,1]):
		table2.rows[a].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		data1 = table2.rows[a].cells[1].paragraphs[0].add_run('%s'%b)
		data1.font.name = 'Myriad Pro'
		data1.font.size = Pt(12)
		table2.rows[a].cells[1].width = Cm(5.5)
		

for c, d in enumerate(infofile2.iloc[:,2]):
		data2 = table2.rows[c].cells[2].paragraphs[0].add_run('%s'%d)
		data2.font.name = 'Cambria'
		data2.font.size = Pt(12)
		table2.rows[c].cells[2].width = Cm(1.8)


for e, f in enumerate(infofile2.iloc[:,3]):
		data3 = table2.rows[e].cells[3].paragraphs[0].add_run('%s'%f)
		data3.font.name = 'Cambria'
		data3.font.size = Pt(12)
		table2.rows[e].cells[3].width = Cm(4.5)

for i, j in enumerate(infofile2.iloc[0,:]):
		table2.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER		
		table2.rows[0].cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER



#################################################################
#Sequencing Result----------------------------------------------#
#################################################################



new_section = document.add_section(1)

SequenceReport = document.add_paragraph()
SequenceReport2 = SequenceReport.add_run('\t\tSequencing Report')
SequenceReport.add_run('\n')
SequenceReport2.font.name = 'Calibri'
SequenceReport2.font.size = Pt(18)
SequenceReport2.font.bold = True
infomation = document.add_paragraph()
infomation2 = infomation.add_run('\t\t Sequencing Result')
infomation2.font.name = 'Calibri'
infomation2.font.size = Pt(14)
infomation2.font.bold = True
infomation.add_run('\n')
infomation3 = infomation.add_run('\t\t The summary table below provides the following overview: read order; index sequence; total \n\t\t  number of bases and reads; percentage of bases in the reads with a phred quality equal or\n\t\t  grater than 30(Q30); average quality score. Individual rows show individual samples. ')
infomation3.font.name = 'Calibri'
infomation3.font.size = Pt(12)
infomation.add_run('\n')


infofile3 = pd.read_csv("./%s_Gumsoo" %ServiceID, sep="\t", header=None, engine="python")
rows3 = len(infofile3.iloc[:,0])
cols3 = len(infofile3.iloc[0,:])
table3 = document.add_table(rows = rows3, cols = cols3)
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
#table3.style = document.styles['Light List Accent 2']
table3.style = document.styles['Medium Grid 2 Accent 3']


for i,j in enumerate(infofile3.iloc[:,0]):
	for x,y in enumerate(infofile3.iloc[i,:]):
		table3.rows[i].cells[x].height = Cm(5)


#for i,j in enumerate(infofile3.iloc[:,0]):
	#for x,y in enumerate(infofile3.iloc[i,:]):
	#	print i, j, x, y
	#	table3.rows[i].cells[0].paragraphs[0].add_run('%s'%y).font.size = Pt(12)
	#	table3.rows[i].cells[x].width = Cm(3.5)
	#	table3.rows[i].cells[x].height = Cm(5)


for i, j in enumerate(infofile3.iloc[:,0]):
		table3.rows[i].cells[0].paragraphs[0].add_run('%s'%j).font.size = Pt(12)
		table3.rows[i].cells[0].width = Cm(3)
		
	


for a, b in enumerate(infofile3.iloc[:,1]):
		table3.rows[a].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
		data1 = table3.rows[a].cells[1].paragraphs[0].add_run('%s'%b)
		data1.font.name = 'Cambria'
		data1.font.size = Pt(12)
		table3.rows[a].cells[1].width = Cm(3.3)
		

for c, d in enumerate(infofile3.iloc[:,2]):
		table3.rows[c].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
		data2 = table3.rows[c].cells[2].paragraphs[0].add_run('%s'%d)
		data2.font.name = 'Cambria'
		data2.font.size = Pt(12)
		table3.rows[c].cells[2].width = Cm(3.3)


for e, f in enumerate(infofile3.iloc[:,3]):
		table3.rows[e].cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		data3 = table3.rows[e].cells[3].paragraphs[0].add_run('%s'%f)
		data3.font.name = 'Cambria'
		data3.font.size = Pt(12)
		table3.rows[e].cells[3].width = Cm(2.8)

for g, h in enumerate(infofile3.iloc[:,4]):
		table3.rows[g].cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		data4 = table3.rows[g].cells[4].paragraphs[0].add_run('%s'%h)
		data4.font.name = 'Cambria'
		data4.font.size = Pt(12)
		table3.rows[g].cells[4].width = Cm(3.0)

for i, j in enumerate(infofile3.iloc[0,:]):
		table3.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER		
		table3.rows[0].cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
#################################################################
#Sequencing graph Yield-----------------------------------------#
#################################################################

new_section = document.add_section(1)

SequenceReport = document.add_paragraph()
SequenceReport2 = SequenceReport.add_run('\t\tSequencing Report')
SequenceReport.add_run('\n')
SequenceReport2.font.name = 'Calibri'
SequenceReport2.font.size = Pt(18)
SequenceReport2.font.bold = True
infomation = document.add_paragraph()
infomation2 = infomation.add_run('\t\t Sequencing Result')
infomation2.font.name = 'Calibri'
infomation2.font.size = Pt(14)
infomation2.font.bold = True
infomation.add_run('\n')
infomation3 = infomation.add_run('\t\t Sequence summary serves as the first step in checking the quality of the raw data. FASTQC was \n\t\t  used to asses the quality of Paired end read sequences. The results of quality control of raw data\n\t\t  is plotted in the two bar plots below. The number of total bases in each sample is shown in the\n\t\t  upper plot, the number of total reads of each sample if shown in the lower plot.')
infomation3.font.name = 'Calibri'
infomation3.font.size = Pt(12)
yieldgraph = document.add_paragraph()
yieldgraph_format = yieldgraph.paragraph_format
center = yieldgraph.add_run()
center.add_picture("./%s_graph_yield.png" %ServiceID, width=Cm(20.03), height=Cm(15.8))
yieldgraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

#################################################################
#Sequencing graph Read------------------------------------------#
#################################################################

new_section = document.add_section(1)

SequenceRead = document.add_paragraph()
SequenceRead2 = SequenceRead.add_run('\t\tSequencing Report')
SequenceRead.add_run('\n')
SequenceRead2.font.name = 'Calibri'
SequenceRead2.font.size = Pt(18)
SequenceRead2.font.bold = True
infomationread = document.add_paragraph()
infomationread2 = infomationread.add_run('\t\t Sequencing Result')
infomationread2.font.name = 'Calibri'
infomationread2.font.size = Pt(14)
infomationread2.font.bold = True
infomationread.add_run('\n')
infomationread3 = infomationread.add_run('\t\t Sequence summary serves as the first step in checking the quality of the raw data. FASTQC was \n\t\t  used to asses the quality of Paired end read sequences. The results of quality control of raw data\n\t\t  is plotted in the two bar plots below. The number of total bases in each sample is shown in the\n\t\t  upper plot, the number of total reads of each sample if shown in the lower plot.')
infomationread3.font.name = 'Calibri'
infomationread3.font.size = Pt(12)
read = document.add_paragraph()
read_format = read.paragraph_format
center = read.add_run()
center.add_picture("./%s_graph_read.png" %ServiceID, width=Cm(20.03), height=Cm(15.8))
read_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

#################################################################
#Download Link Page---------------------------------------------#
#################################################################
new_section = document.add_section(1)

with open("./%s_Link" %ServiceID, mode='r') as f:
	 Read = f.read()
	 Link_ls = Read.split()

DownloadLink_PAGE = document.add_paragraph()
DownloadLink = DownloadLink_PAGE.add_run('\t\tDownload Link\n')
DownloadLink.font.name = 'Calibri'
DownloadLink.font.size = Pt(18)
DownloadLink.font.bold = True
DownloadLink2 = DownloadLink_PAGE.add_run('\t\tRawdata\n')
DownloadLink2.font.name = 'Calibri'
DownloadLink2.font.size = Pt(14)
DownloadLink2.font.bold = True
ID = DownloadLink_PAGE.add_run('\t\tDownload ID')
ID.font.name = 'Helvetica'
ID.font.size = Pt(10.5)
ID.font.bold = True
ID2 = DownloadLink_PAGE.add_run('\t\tuser1\n')
ID2.font.name = 'Helvetica'
ID2.font.size = Pt(11)
ID2.font.color.rgb = RGBColor(30, 144, 255)
PW = DownloadLink_PAGE.add_run('\t\tDownload PW')
PW.font.name = 'Helvetica'
PW.font.size = Pt(10.5)
PW.font.bold = True
PW2 = DownloadLink_PAGE.add_run('\t\tdnalink1234%\n')
PW2.font.name = 'Helvetica'
PW2.font.size = Pt(11)
PW2.font.color.rgb = RGBColor(30, 144, 255)

## 파일 링크 입력
for Line in Link_ls:
	FileName = Line.split("/")[-1]
	HL(DownloadLink_PAGE, Line, "\t\t%s\n" % FileName, "800f7", True)

#document.add_page_break()
document.save('%s_Sequecing_Report.docx' % ServiceID)
