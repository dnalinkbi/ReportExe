import docx
from tkinter import *

# Create the window
root = Tk()
root.title("My Word Creator")

# Create the labels and text entry boxes
label_1 = Label(root, text="Enter the title of your document:")
entry_1 = Entry(root)
label_2 = Label(root, text="Enter the body of your document:")
entry_2 = Text(root, height=10, width=50)

# Place the widgets in the window
label_1.grid(row=0, column=0, sticky=E)
entry_1.grid(row=0, column=1)
label_2.grid(row=1, column=0, sticky=E)
entry_2.grid(row=1, column=1)

# Create the function to write the document
def create_document():
    document = docx.Document()
    document.add_heading(entry_1.get(), 0)
    document.add_paragraph(entry_2.get("1.0", END))
    document.save('my_document.docx')

# Create the button
button = Button(root, text="Create Document", command=create_document)
button.grid(row=2, columnspan=2)

# Run the window
root.mainloop()